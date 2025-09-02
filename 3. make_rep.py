import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import re
import matplotlib.pyplot as plt
from docx.shared import Inches
from collections import defaultdict
import matplotlib.font_manager as fm

# Установка шрифта для matplotlib
def setup_matplotlib_font():
    try:
        # Ищем файл шрифта в каталоге PFCentroSansPro
        font_dir = "PFCentroSansPro"
        if os.path.exists(font_dir):
            # Ищем файлы шрифтов в каталоге
            font_files = []
            for file in os.listdir(font_dir):
                if file.lower().endswith(('.ttf', '.otf')):
                    font_files.append(os.path.join(font_dir, file))
            
            if font_files:
                # Используем первый найденный файл шрифта
                prop = fm.FontProperties(fname=font_files[0])
                plt.rcParams['font.family'] = prop.get_name()
                print(f"Установлен шрифт: {prop.get_name()}")
            else:
                print("В каталоге PFCentroSansPro не найдено файлов шрифтов. Используется стандартный шрифт.")
        else:
            print("Каталог PFCentroSansPro не найден. Используется стандартный шрифт.")
    except Exception as e:
        print(f"Ошибка при настройке шрифта: {e}")

def set_cell_text(cell, text, font_name='PF Centro Sans Pro', font_size=Pt(14)):
    """Устанавливает текст в ячейке с указанным шрифтом и размером"""
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text) if text is not None else '')
    run.font.name = font_name
    run.font.size = font_size
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def get_russian_month():
    """Возвращает предыдущий месяц на русском с заглавной буквы"""
    months = {
        1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель",
        5: "Май", 6: "Июнь", 7: "Июль", 8: "Август",
        9: "Сентябрь", 10: "Октябрь", 11: "Ноябрь", 12: "Декабрь"
    }
    
    # Получаем текущую дату
    today = datetime.now()
    
    # Вычисляем первый день текущего месяца
    first_day_of_current_month = today.replace(day=1)
    
    # Вычитаем один день, чтобы получить последний день предыдущего месяца
    last_day_of_previous_month = first_day_of_current_month - timedelta(days=1)
    
    # Получаем номер предыдущего месяца
    previous_month = last_day_of_previous_month.month
    
    return months.get(previous_month, "")

def replace_placeholder(doc, placeholder, replacement):
    """Заменяет плейсхолдеры в документе, включая титульную страницу"""
    replacement_str = str(replacement) if replacement is not None else ""
    
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement_str)
            for run in paragraph.runs:
                run.font.name = 'PF Centro Sans Pro'
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, replacement_str)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'PF Centro Sans Pro'

def find_mitre_table(doc):
    """Находит таблицу с тактиками и техниками MITRE ATT&CK по контексту"""
    table_title = "Таблица 2. Тактики и техники"
    
    for paragraph in doc.paragraphs:
        if table_title in paragraph.text:
            for element in paragraph._element.xpath('./following::*'):
                if element.tag.endswith('tbl'):
                    for table in doc.tables:
                        if table._element == element:
                            return table
    return None

def fill_mitre_table(doc, techniques, technique_to_tactic):
    """Заполняет таблицу с тактиками и техниками MITRE ATT&CK"""
    target_table = find_mitre_table(doc)
    
    if not target_table:
        print("Таблица с тактиками и техниками MITRE ATT&CK не найдена в шаблоне")
        return
    
    for row in list(target_table.rows)[1:]:
        target_table._tbl.remove(row._tr)
    
    for technique in techniques:
        tactic = technique_to_tactic.get(str(technique).strip() if technique else "", "Неизвестная тактика")
        
        new_row = target_table.add_row()
        while len(new_row.cells) < 2:
            new_row.add_cell()
        
        set_cell_text(new_row.cells[0], tactic)
        set_cell_text(new_row.cells[1], technique)

def create_tactics_chart(techniques, technique_to_tactic, output_dir, company_name):
    """Создаёт столбчатую диаграмму распределения тактик MITRE ATT&CK"""
    if not techniques:
        print("Нет данных для построения диаграммы")
        return None
    
    tactic_count = defaultdict(int)
    for tech in techniques:
        tactic = technique_to_tactic.get(str(tech).strip(), "Неизвестная тактика")
        tactic_count[tactic] += 1
    
    if not tactic_count:
        return None
    
    setup_matplotlib_font()
    plt.figure(figsize=(12, 8))
    tactics = list(tactic_count.keys())
    counts = list(tactic_count.values())
    
    bars = plt.bar(tactics, counts, color='skyblue')
    plt.xlabel('Тактики MITRE ATT&CK', fontsize=12)
    plt.ylabel('Количество правил', fontsize=12)
    plt.title('Распределение правил по тактикам MITRE ATT&CK', fontsize=14)
    plt.xticks(rotation=45, ha='right')
    
    for bar, count in zip(bars, counts):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                str(count), ha='center', va='bottom')
    
    plt.tight_layout()
    
    # Создаем безопасное имя файла для компании
    safe_company_name = re.sub(r'[\\/*?:"<>|]', "_", company_name)
    chart_path = os.path.join(output_dir, f"{safe_company_name}_tactics_chart.png")
    
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path

def replace_chart_placeholder(doc, chart_path):
    """Заменяет плейсхолдер {chart} на диаграмму в документе"""
    # Ищем плейсхолдер {chart} в параграфах
    for paragraph in doc.paragraphs:
        if "{chart}" in paragraph.text:
            # Очищаем параграф и вставляем изображение
            paragraph.text = ""
            run = paragraph.add_run()
            run.add_picture(chart_path, width=Inches(6))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return True
    
    # Ищем плейсхолдер {chart} в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{chart}" in cell.text:
                    # Очищаем ячейку и вставляем изображение
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(chart_path, width=Inches(6))
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    return True
    
    print("Плейсхолдер {chart} не найден в документе")
    return False

def format_period(period_str):
    """Преобразует период из формата '2025-05-01T00:00:00+03:00 - 2025-05-31T00:00:00+03:00' в '01.05.2025 - 31.05.2025'"""
    if not period_str or pd.isna(period_str):
        return "[период не указан]"
    
    try:
        period_str = str(period_str)
        dates = re.findall(r'(\d{4}-\d{2}-\d{2})', period_str)
        if len(dates) == 2:
            start_date = datetime.strptime(dates[0], "%Y-%m-%d").strftime("%d.%m.%Y")
            end_date = datetime.strptime(dates[1], "%Y-%m-%d").strftime("%d.%m.%Y")
            return f"{start_date} - {end_date}"
        return period_str
    except Exception as e:
        print(f"Ошибка форматирования периода: {e}")
        return str(period_str)

def find_sources_table(doc):
    """Находит таблицу с затронутыми источниками по контексту"""
    table_title = "Таблица 1. Имена затронутых источников"
    
    for paragraph in doc.paragraphs:
        if table_title in paragraph.text:
            for element in paragraph._element.xpath('./following::*'):
                if element.tag.endswith('tbl'):
                    for table in doc.tables:
                        if table._element == element:
                            return table
    return None

def fill_sources_table(doc, assets):
    """Заполняет таблицу с затронутыми источниками"""
    target_table = find_sources_table(doc)
    
    if not target_table:
        print("Таблица с затронутыми источниками не найдена в шаблоне")
        return
    
    for row in list(target_table.rows)[1:]:
        target_table._tbl.remove(row._tr)
    
    rows_needed = min(35, (len(assets) + 3) // 4)
    
    for _ in range(rows_needed):
        new_row = target_table.add_row()
        
    for i in range(rows_needed):
        row = target_table.rows[i + 1]
        for j in range(4):
            idx = i * 4 + j
            if idx < len(assets):
                if j*2 + 1 < len(row.cells):
                    set_cell_text(row.cells[j*2], str(assets[idx][0]) if assets[idx][0] is not None else "")
                    set_cell_text(row.cells[j*2 + 1], str(assets[idx][1]) if assets[idx][1] is not None else "")
            else:
                if j*2 + 1 < len(row.cells):
                    set_cell_text(row.cells[j*2], '')
                    set_cell_text(row.cells[j*2 + 1], '')

def process_excel_file(excel_path, template_path, output_dir):
    """Обрабатывает один Excel файл и генерирует отчет"""
    try:
        df_1_1 = pd.read_excel(excel_path, sheet_name="1-1", header=None, nrows=1)
        company_name = str(df_1_1.iloc[0, 2]) if len(df_1_1.columns) > 2 and not pd.isna(df_1_1.iloc[0, 2]) else "[Название организации]"
        events_count = str(df_1_1.iloc[0, 6]) if len(df_1_1.columns) > 6 and not pd.isna(df_1_1.iloc[0, 6]) else "[неизвестно]"
        period = format_period(df_1_1.iloc[0, 1] if len(df_1_1.columns) > 1 else None)
        
        df_1_2 = pd.read_excel(excel_path, sheet_name="1-2", header=None, nrows=1)
        alerts_count = str(df_1_2.iloc[0, 7]) if len(df_1_2.columns) > 7 and not pd.isna(df_1_2.iloc[0, 7]) else "[неизвестно]"
        
        try:
            df_1_5 = pd.read_excel(excel_path, sheet_name="1-5", header=None, usecols="E", skiprows=2)
            assets = [(i+1, str(val)) for i, val in enumerate(df_1_5.iloc[:, 0]) if not df_1_5.empty and not pd.isna(val)] if not df_1_5.empty else []
        except Exception as e:
            print(f"Ошибка при чтении листа 1-5: {e}")
            assets = []
        
        try:
            df_1_6 = pd.read_excel(excel_path, sheet_name="1-6", header=None, usecols="A", skiprows=2)
            techniques = [str(val).strip() for val in df_1_6.iloc[:, 0] if not df_1_6.empty and not pd.isna(val)] if not df_1_6.empty else []
        except Exception as e:
            print(f"Ошибка при чтении листа 1-6: {e}")
            techniques = []
        
    except Exception as e:
        print(f"Ошибка при чтении Excel файла {excel_path}: {e}")
        return

    try:
        mitre_mapping = pd.read_excel("rules.xlsx", sheet_name="Sheet1")
        technique_to_tactic = dict(zip(
            mitre_mapping['Original_Rule'].astype(str).str.strip(),
            mitre_mapping['MITRE_Tactic'].astype(str)
        ))
    except Exception as e:
        print(f"Ошибка при чтении файла сопоставления MITRE: {e}")
        technique_to_tactic = {}

    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Ошибка при загрузке шаблона DOCX: {e}")
        return

    replace_data = {
        "{предпр}": company_name,
        "{Месяц}": get_russian_month(),
        "{месяц}": get_russian_month(),
        "{год}": str(datetime.now().year),
        "{число}": events_count,
        "{алертов}": alerts_count,
        "{ист}": str(len(assets)),
        "{период}": period
    }
    
    for placeholder, value in replace_data.items():
        replace_placeholder(doc, placeholder, value)

    fill_sources_table(doc, assets)
    fill_mitre_table(doc, techniques, technique_to_tactic)
    
    # Создаем и вставляем диаграмму вместо плейсхолдера {chart}
    chart_path = create_tactics_chart(techniques, technique_to_tactic, output_dir, company_name)
    if chart_path:
        replace_chart_placeholder(doc, chart_path)
        # Удаляем временный файл диаграммы
        try:
            os.remove(chart_path)
        except Exception as e:
            print(f"Ошибка при удалении временного файла диаграммы: {e}")
   
    base_name = os.path.basename(excel_path)
    report_name = os.path.splitext(base_name)[0] + "_report.docx"
    output_path = os.path.join(output_dir, report_name)

    try:
        doc.save(output_path)
        print(f"Отчет успешно сгенерирован: {output_path}")
    except Exception as e:
        print(f"Ошибка при сохранении отчета: {e}")

def generate_reports():
    if not os.path.exists("output"):
        print("Каталог 'output' не существует")
        return
    
    if not os.path.exists("template.docx"):
        print("Файл шаблона 'template.docx' не найден")
        return
    
    os.makedirs("reports", exist_ok=True)
    
    excel_files = [f for f in os.listdir("output") if f.lower().endswith(('.xlsx', '.xls'))]
    
    if not excel_files:
        print("В каталоге 'output' не найдено Excel файлов")
        return
    
    for excel_file in excel_files:
        excel_path = os.path.join("output", excel_file)
        print(f"Обработка файла: {excel_path}")
        process_excel_file(excel_path, "template.docx", "reports")

if __name__ == "__main__":
    generate_reports()