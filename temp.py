import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import matplotlib.pyplot as plt
from docx.shared import Inches
from collections import defaultdict

def set_cell_text(cell, text, font_name='PF Centro Sans Pro', font_size=Pt(12)):
    """Устанавливает текст в ячейке с указанным шрифтом и размером"""
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text) if text is not None else '')
    run.font.name = font_name
    run.font.size = font_size
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def get_russian_month():
    """Возвращает предыдущий месяц на русском с заглавной буквы"""
    today = datetime.now()
    first_day = today.replace(day=1)
    prev_month = first_day - timedelta(days=1)
    
    months = {
        1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель",
        5: "Май", 6: "Июнь", 7: "Июль", 8: "Август",
        9: "Сентябрь", 10: "Октябрь", 11: "Ноябрь", 12: "Декабрь"
    }
    return months.get(prev_month.month, "")

def set_font_size(doc, size=Pt(12)):
    """Устанавливает размер шрифта для всего документа"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = size
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = size

def replace_placeholder(doc, placeholder, replacement):
    """Заменяет плейсхолдеры в документе"""
    replacement_str = str(replacement) if replacement is not None else ""
    
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement_str)
            for run in paragraph.runs:
                run.font.name = 'PF Centro Sans Pro'
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, replacement_str)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'PF Centro Sans Pro'

def find_mitre_table(doc):
    """Находит таблицу с тактиками и техниками MITRE ATT&CK"""
    table_title = "Таблица 2. Тактики и техники"
    
    for paragraph in doc.paragraphs:
        if table_title in paragraph.text:
            for element in paragraph._element.xpath('./following::*'):
                if element.tag.endswith('tbl'):
                    for table in doc.tables:
                        if table._element == element:
                            return table
    return None

def fill_mitre_table(doc, techniques):
    """Заполняет таблицу с тактиками и техниками MITRE ATT&CK"""
    target_table = find_mitre_table(doc)
    
    if not target_table:
        print("Таблица MITRE ATT&CK не найдена")
        return
    
    for row in list(target_table.rows)[1:]:
        target_table._tbl.remove(row._tr)
    
    try:
        mitre_mapping = pd.read_excel("123.xlsx", sheet_name="Processed Data")
        technique_to_tactic = dict(zip(
            mitre_mapping['Original'].astype(str),
            mitre_mapping['MITRE Match'].astype(str)
        ))
    except Exception as e:
        print(f"Ошибка чтения файла MITRE: {e}")
        technique_to_tactic = {}
    
    for technique in techniques:
        tactic = technique_to_tactic.get(str(technique).strip() if technique else "", "Неизвестная тактика")
        
        new_row = target_table.add_row()
        while len(new_row.cells) < 2:
            new_row.add_cell()
        
        set_cell_text(new_row.cells[0], tactic)
        set_cell_text(new_row.cells[1], technique)

def format_period(period_str):
    """Форматирует период отображения"""
    if not period_str or pd.isna(period_str):
        return "[период не указан]"
    
    try:
        period_str = str(period_str)
        dates = re.findall(r'(\d{4}-\d{2}-\d{2})', period_str)
        if len(dates) == 2:
            start_date = datetime.strptime(dates[0], "%Y-%m-%d").strftime("%d.%m.%Y")
            end_date = datetime.strptime(dates[1], "%Y-%m-%d").strftime("%d.%m.%Y")
            return f"{start_date} - {end_date}"
        return period_str
    except Exception as e:
        print(f"Ошибка форматирования периода: {e}")
        return str(period_str)

def find_sources_table(doc):
    """Находит таблицу с затронутыми источниками"""
    table_title = "Таблица 1. Имена затронутых источников"
    
    for paragraph in doc.paragraphs:
        if table_title in paragraph.text:
            for element in paragraph._element.xpath('./following::*'):
                if element.tag.endswith('tbl'):
                    for table in doc.tables:
                        if table._element == element:
                            return table
    return None

def fill_sources_table(doc, assets):
    """Заполняет таблицу с затронутыми источниками"""
    target_table = find_sources_table(doc)
    
    if not target_table:
        print("Таблица источников не найдена")
        return
    
    for row in list(target_table.rows)[1:]:
        target_table._tbl.remove(row._tr)
    
    rows_needed = min(35, (len(assets) + 3) // 4)
    
    for _ in range(rows_needed):
        new_row = target_table.add_row()
        
    for i in range(rows_needed):
        row = target_table.rows[i + 1]
        for j in range(4):
            idx = i * 4 + j
            if idx < len(assets):
                if j*2 + 1 < len(row.cells):
                    set_cell_text(row.cells[j*2], str(assets[idx][0]) if assets[idx][0] is not None else "")
                    set_cell_text(row.cells[j*2 + 1], str(assets[idx][1]) if assets[idx][1] is not None else "")
                else:
                    print(f"Недостаточно ячеек в строке {i+1}")
            else:
                if j*2 + 1 < len(row.cells):
                    set_cell_text(row.cells[j*2], '')
                    set_cell_text(row.cells[j*2 + 1], '')
                    
def create_mitre_chart(techniques, excel_path):
    """Создает диаграмму распределения тактик MITRE"""
    if not techniques:
        return None
        
    try:
        mitre_mapping = pd.read_excel("123.xlsx", sheet_name="Processed Data")
        technique_to_tactic = dict(zip(
            mitre_mapping['Original'].astype(str),
            mitre_mapping['MITRE Match'].astype(str)
        ))
    except Exception as e:
        print(f"Ошибка чтения файла MITRE: {e}")
        return None
    
    tactic_counts = defaultdict(int)
    for technique in techniques:
        technique_str = str(technique).strip() if technique else ""
        tactic = technique_to_tactic.get(technique_str, "Неизвестная тактика")
        if tactic != "Неизвестная тактика":
            tactic_counts[tactic] += 1
    
    if not tactic_counts:
        return None
    
    sorted_tactics = sorted(tactic_counts.items(), key=lambda x: x[1], reverse=True)
    tactics, counts = zip(*sorted_tactics)
    
    plt.figure(figsize=(12, 6))
    bars = plt.bar(tactics, counts, color='#4472C4')
    plt.xlabel('Тактики MITRE ATT&CK')
    plt.ylabel('Количество техник')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height,
                 f'{int(height)}',
                 ha='center', va='bottom')
    
    chart_path = os.path.join(os.path.dirname(excel_path), "mitre_chart.png")
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path

def find_page3_chart(doc):
    """Находит существующую диаграмму на 3 странице"""
    # Поиск по заголовку страницы
    page3_title = "3. Результаты анализа угроз"
    page3_index = -1
    
    # Находим начало 3 страницы
    for i, paragraph in enumerate(doc.paragraphs):
        if page3_title in paragraph.text:
            page3_index = i
            break
    
    if page3_index == -1:
        return None
    
    # Ищем первый графический элемент на странице
    for i in range(page3_index, min(page3_index + 50, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        for run in para.runs:
            if run._element.xpath('.//pic:pic'):
                return para
    return None

def replace_existing_chart(doc, chart_path):
    """Заменяет существующую диаграмму на 3 странице новой"""
    chart_para = find_page3_chart(doc)
    
    if chart_para is None:
        print("Диаграмма на 3 странице не найдена, вставляем в конец")
        return insert_chart_fallback(doc, chart_path)
    
    # Сохраняем выравнивание
    alignment = chart_para.alignment
    
    # Очищаем параграф
    chart_para.clear()
    
    # Вставляем новую диаграмму
    run = chart_para.add_run()
    try:
        run.add_picture(chart_path, width=Inches(6.5))
        chart_para.alignment = alignment
        return True
    except Exception as e:
        print(f"Ошибка вставки изображения: {e}")
        return insert_chart_fallback(doc, chart_path)

def insert_chart_fallback(doc, chart_path):
    """Вставляет диаграмму после таблицы MITRE (резервный вариант)"""
    mitre_table = find_mitre_table(doc)
    if not mitre_table:
        print("Не удалось найти таблицу MITRE")
        return False
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    
    try:
        run.add_picture(chart_path, width=Inches(6.5))
    except Exception as e:
        print(f"Ошибка вставки изображения: {e}")
        return False
    
    caption = doc.add_paragraph("Рисунок 1. Распределение техник по тактикам MITRE ATT&CK")
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in caption.runs:
        run.font.name = 'PF Centro Sans Pro'
        run.font.size = Pt(12)
    
    return True

def process_excel_file(excel_path, template_path, output_dir):
    """Обрабатывает Excel файл и генерирует отчет"""
    try:
        df_1_1 = pd.read_excel(excel_path, sheet_name="1-1", header=None, nrows=1)
        company_name = str(df_1_1.iloc[0, 2]) if len(df_1_1.columns) > 2 and not pd.isna(df_1_1.iloc[0, 2]) else "[Название организации]"
        events_count = str(df_1_1.iloc[0, 6]) if len(df_1_1.columns) > 6 and not pd.isna(df_1_1.iloc[0, 6]) else "[неизвестно]"
        period = format_period(df_1_1.iloc[0, 1] if len(df_1_1.columns) > 1 else None)
        
        df_1_2 = pd.read_excel(excel_path, sheet_name="1-2", header=None, nrows=1)
        alerts_count = str(df_1_2.iloc[0, 7]) if len(df_1_2.columns) > 7 and not pd.isna(df_1_2.iloc[0, 7]) else "[неизвестно]"
        
        try:
            df_1_5 = pd.read_excel(excel_path, sheet_name="1-5", header=None, usecols="E", skiprows=2)
            assets = [(i+1, str(val)) for i, val in enumerate(df_1_5.iloc[:, 0]) if not df_1_5.empty and not pd.isna(val)] if not df_1_5.empty else []
        except Exception as e:
            print(f"Ошибка чтения листа 1-5: {e}")
            assets = []
        
        try:
            df_1_6 = pd.read_excel(excel_path, sheet_name="1-6", header=None, usecols="A", skiprows=2)
            techniques = [str(val).strip() for val in df_1_6.iloc[:, 0] if not df_1_6.empty and not pd.isna(val)] if not df_1_6.empty else []
        except Exception as e:
            print(f"Ошибка чтения листа 1-6: {e}")
            techniques = []
        
    except Exception as e:
        print(f"Ошибка чтения Excel: {e}")
        return

    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Ошибка загрузки шаблона: {e}")
        return

    replace_data = {
        "{предпр}": company_name,
        "{Месяц}": get_russian_month(),
        "{месяц}": get_russian_month(),
        "{год}": str(datetime.now().year),
        "{число}": events_count,
        "{алертов}": alerts_count,
        "{ист}": str(len(assets)),
        "{период}": period
    }
    
    for placeholder, value in replace_data.items():
        replace_placeholder(doc, placeholder, value)

    fill_sources_table(doc, assets)
    fill_mitre_table(doc, techniques)
    
    chart_path = create_mitre_chart(techniques, excel_path)
    if chart_path:
        replace_existing_chart(doc, chart_path)
        try:
            os.remove(chart_path)
        except Exception as e:
            print(f"Ошибка удаления временного файла: {e}")
    
    # Установка размера шрифта 12 для всего документа
    set_font_size(doc, Pt(12))

    report_name = os.path.splitext(os.path.basename(excel_path))[0] + "_report.docx"
    output_path = os.path.join(output_dir, report_name)

    try:
        doc.save(output_path)
        print(f"Отчет сохранен: {output_path}")
    except Exception as e:
        print(f"Ошибка сохранения отчета: {e}")

def generate_reports():
    if not os.path.exists("output"):
        os.makedirs("output")
        print("Создан каталог 'output'")
    
    if not os.path.exists("template.docx"):
        print("Файл шаблона 'template.docx' не найден")
        return
    
    os.makedirs("reports", exist_ok=True)
    
    excel_files = [f for f in os.listdir("output") if f.lower().endswith(('.xlsx', '.xls'))]
    
    if not excel_files:
        print("В каталоге 'output' нет Excel файлов")
        return
    
    for excel_file in excel_files:
        excel_path = os.path.join("output", excel_file)
        print(f"Обработка: {excel_path}")
        process_excel_file(excel_path, "template.docx", "reports")

if __name__ == "__main__":
    generate_reports()